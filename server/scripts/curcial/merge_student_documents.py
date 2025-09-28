#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
学生论文材料合并脚本
功能：
1. 读取Excel中的学生信息
2. 为每个学生更新封面信息
3. 将Word文档转换为PDF
4. 按文件名数字顺序合并PDF
5. 创建带导航的最终PDF
"""

import os
import re
import pandas as pd
import fitz  # PyMuPDF
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import subprocess
import sys
from pathlib import Path
import json

# 设置环境变量确保UTF-8编码
os.environ['PYTHONIOENCODING'] = 'utf-8'

class StudentDocumentMerger:
    def __init__(self, student_folders, output_dir, base_dir=None, excel_file=None):
        # base_dir为curcial目录
        self.base_dir = Path(base_dir) if base_dir else Path(__file__).parent
        self.student_folders = [Path(f) for f in student_folders]
        if excel_file:
            self.excel_file = Path(excel_file)
        else:
            self.excel_file = self.base_dir / "学号&姓名&专业.xlsx"
        self.cover_template = self.base_dir / "通用封面.docx"
        self.output_dir = Path(output_dir)
        self.student_info = self.load_student_info()
        self.output_dir.mkdir(exist_ok=True, parents=True)
        
        # 检查封面模板文件是否存在
        if not self.cover_template.exists():
            print(f"警告: 封面模板文件不存在: {self.cover_template}")
            print(f"请确保在 {self.base_dir} 目录下有 '通用封面.docx' 文件")

    def load_student_info(self):
        try:
            # 检查Excel文件是否存在
            if not self.excel_file.exists():
                print(f"错误: Excel文件不存在: {self.excel_file}")
                return {}
            
            print(f"正在读取Excel文件: {self.excel_file}")
            df = pd.read_excel(self.excel_file, dtype={'学号': str})
            print(f"成功读取学生信息，共 {len(df)} 名学生")
            print(f"Excel列名: {list(df.columns)}")
            print(f"前几行数据:\n{df.head()}")
            
            # 检查必要的列是否存在
            required_columns = ['姓名', '学号', '专业']
            missing_columns = [col for col in required_columns if col not in df.columns]
            if missing_columns:
                print(f"错误: Excel文件缺少必要的列: {missing_columns}")
                return {}
            
            return df.set_index('姓名').to_dict('index')
        except Exception as e:
            print(f"读取Excel文件失败: {e}")
            import traceback
            traceback.print_exc()
            return {}

    def update_cover_for_student(self, student_name):
        if student_name not in self.student_info:
            print(f"警告: 在Excel中找不到学生 {student_name} 的信息")
            return None
        student_data = self.student_info[student_name]
        student_num = str(student_data['学号'])
        major = student_data['专业']
        try:
            # 检查封面模板文件是否存在
            if not self.cover_template.exists():
                print(f"错误: 封面模板文件不存在: {self.cover_template}")
                return None
            print(f"正在为 {student_name} 创建封面，使用模板: {self.cover_template}")
            doc = Document(self.cover_template)
            for paragraph in doc.paragraphs:
                full_text = paragraph.text
                if '专业' in full_text or '学号' in full_text or '姓名' in full_text:
                    new_text = full_text
                    new_text = re.sub(r'专业[：:]\s*[\w\s]*', f'专业：{major}', new_text)
                    new_text = re.sub(r'学号[：:]\s*[\w\d]*', f'学号：{student_num}', new_text)
                    new_text = re.sub(r'姓名[：:]\s*[\w\s]*', f'姓名：{student_name}', new_text)
                    if new_text != full_text:
                        paragraph.clear()
                        # 设置段落对齐方式为左对齐
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        # 清除所有缩进和边距设置
                        paragraph.paragraph_format.left_indent = 0
                        paragraph.paragraph_format.right_indent = 0
                        paragraph.paragraph_format.first_line_indent = 0
                        paragraph.paragraph_format.hanging_indent = 0
                        # 设置段落边距为0
                        paragraph.paragraph_format.space_before = 0
                        paragraph.paragraph_format.space_after = 0
                        # 强制设置左边距
                        paragraph.paragraph_format.left_indent = 0  # 负缩进，向左移动
                        paragraph.add_run(new_text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            full_text = paragraph.text
                            if '专业' in full_text or '学号' in full_text or '姓名' in full_text:
                                new_text = full_text
                                new_text = re.sub(r'专业[：:]\s*[\w\s]*', f'专业：{major}', new_text)
                                new_text = re.sub(r'学号[：:]\s*[\w\d]*', f'学号：{student_num}', new_text)
                                new_text = re.sub(r'姓名[：:]\s*[\w\s]*', f'姓名：{student_name}', new_text)
                                if new_text != full_text:
                                    paragraph.clear()
                                    # 设置段落对齐方式为左对齐
                                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                                    # 清除所有缩进和边距设置
                                    paragraph.paragraph_format.left_indent = 0
                                    paragraph.paragraph_format.right_indent = 0
                                    paragraph.paragraph_format.first_line_indent = 0
                                    paragraph.paragraph_format.hanging_indent = 0
                                    # 设置段落边距为0
                                    paragraph.paragraph_format.space_before = 0
                                    paragraph.paragraph_format.space_after = 0
                                    # 强制设置左边距
                                    paragraph.paragraph_format.left_indent = 0 # 负缩进，向左移动
                                    paragraph.add_run(new_text)
            # 确保封面文件名使用正确的UTF-8编码
            safe_student_name = self.sanitize_filename(student_name)
            cover_filename = f"封面_{safe_student_name}.docx"
            cover_path = self.safe_path_join(self.output_dir, cover_filename)
            
            print(f"正在保存封面到: {cover_path}")
            print(f"封面文件名: {cover_filename}")
            
            # 确保输出目录存在
            cover_path.parent.mkdir(parents=True, exist_ok=True)
            
            doc.save(cover_path)
            
            # 验证封面文件是否成功保存
            if cover_path.exists():
                file_size = cover_path.stat().st_size
                print(f"✓ 已为 {student_name} 创建个人封面: {cover_path}")
                print(f"  封面文件大小: {file_size} 字节")
                if file_size == 0:
                    print("  警告: 封面文件大小为0，可能保存失败")
                return cover_path
            else:
                print(f"✗ 封面文件保存失败: {cover_path}")
                return None
        except Exception as e:
            print(f"更新封面失败 {student_name}: {e}")
            return None

    def word_to_pdf(self, word_file, output_dir):
        try:
            # 检查输入文件是否存在
            if not Path(word_file).exists():
                print(f"错误: 输入文件不存在: {word_file}")
                return None
            
            # 检查soffice命令是否可用
            try:
                version_result = subprocess.run(['soffice', '--version'], capture_output=True, check=True, text=True)
                print(f"LibreOffice版本: {version_result.stdout.strip()}")
            except (subprocess.CalledProcessError, FileNotFoundError) as e:
                print(f"错误: 未找到LibreOffice的soffice命令: {e}")
                print("请确保已安装LibreOffice并添加到PATH环境变量中")
                print("安装方法:")
                print("  macOS: brew install --cask libreoffice")
                print("  Ubuntu: sudo apt install libreoffice")
                print("  Windows: 从官网下载安装包")
                return None
            
            cmd = [
                'soffice', '--headless', '--convert-to', 'pdf',
                '--outdir', str(output_dir), str(word_file)
            ]
            print(f"正在转换: {word_file} -> PDF")
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
            
            if result.returncode == 0:
                pdf_name = Path(word_file).stem + '.pdf'
                pdf_path = Path(output_dir) / pdf_name
                if pdf_path.exists():
                    print(f"成功转换: {word_file} -> {pdf_name}")
                    print(f"转换后的PDF文件大小: {pdf_path.stat().st_size} 字节")
                    return pdf_path
                else:
                    print(f"转换后找不到PDF文件: {pdf_path}")
                    print(f"输出目录内容: {list(Path(output_dir).iterdir())}")
                    # 尝试查找可能的PDF文件
                    for file in Path(output_dir).iterdir():
                        if file.suffix.lower() == '.pdf':
                            print(f"发现可能的PDF文件: {file}")
                    return None
            else:
                print(f"soffice转换失败 (返回码: {result.returncode})")
                print(f"错误输出: {result.stderr}")
                print(f"标准输出: {result.stdout}")
                return None
        except subprocess.TimeoutExpired:
            print(f"转换超时: {word_file}")
            return None
        except FileNotFoundError:
            print("错误: 未找到soffice命令，请安装LibreOffice")
            return None
        except Exception as e:
            print(f"转换出错: {e}")
            import traceback
            traceback.print_exc()
            return None

    def extract_file_number(self, filename):
        """提取文件名开头的数字用于排序"""
        match = re.match(r'^(\d+)', filename)
        if match:
            return int(match.group(1))
        else:
            # 对于没有数字开头的文件，根据文件名进行排序
            # 将文件名转换为小写并按字母顺序排序
            return 1000 + hash(filename) % 1000
    
    def ensure_utf8_string(self, text):
        """确保字符串是UTF-8编码的"""
        if isinstance(text, bytes):
            try:
                return text.decode('utf-8')
            except UnicodeDecodeError:
                return text.decode('gbk', errors='ignore')
        return str(text)
    
    def sanitize_filename(self, filename):
        """清理文件名，移除或替换不安全的字符，保留中文字符"""
        if not filename:
            return "unnamed_file"
        
        # 确保是字符串
        filename = str(filename)
        
        # 移除或替换不安全的字符，但保留中文字符
        unsafe_chars = ['\x00', '\x01', '\x02', '\x03', '\x04', '\x05', '\x06', '\x07', 
                       '\x08', '\x0b', '\x0c', '\x0e', '\x0f', '\x10', '\x11', '\x12', 
                       '\x13', '\x14', '\x15', '\x16', '\x17', '\x18', '\x19', '\x1a', 
                       '\x1b', '\x1c', '\x1d', '\x1e', '\x1f', '<', '>', ':', '"', '|', 
                       '?', '*', '\\', '/']
        
        for char in unsafe_chars:
            filename = filename.replace(char, '_')
        
        # 移除前后空格和点
        filename = filename.strip(' .')
        
        # 如果文件名为空，使用默认名称
        if not filename:
            filename = "unnamed_file"
        
        return filename
    
    def safe_path_join(self, *parts):
        """安全地连接路径，处理中文编码问题"""
        try:
            # 将所有部分转换为字符串并确保UTF-8编码
            safe_parts = []
            for part in parts:
                if isinstance(part, bytes):
                    part = part.decode('utf-8', errors='ignore')
                else:
                    part = str(part)
                # 清理路径中的特殊字符
                part = part.replace('\x00', '').replace('\x01', '').replace('\x02', '')
                safe_parts.append(part)
            
            # 使用Path.joinpath方法
            result = Path(safe_parts[0])
            for part in safe_parts[1:]:
                result = result / part
            
            # 确保路径是绝对路径
            result = result.resolve()
            return result
        except Exception as e:
            print(f"路径连接失败: {e}")
            # 回退到简单字符串连接
            try:
                result = Path(str(parts[0])).joinpath(*[str(p) for p in parts[1:]])
                return result.resolve()
            except Exception as e2:
                print(f"回退路径连接也失败: {e2}")
                # 最后回退：使用当前目录
                return Path.cwd() / "temp_output.pdf"

    def get_student_files(self, student_folder):
        files = []
        try:
            student_path = Path(student_folder)
            print(f"正在扫描学生文件夹: {student_path}")
            print(f"文件夹是否存在: {student_path.exists()}")
            
            if not student_path.exists():
                print(f"错误: 学生文件夹不存在: {student_path}")
                return files
            
            for file_path in student_path.iterdir():
                if file_path.is_file() and not file_path.name.startswith('.'):
                    # 确保文件名正确编码
                    try:
                        # 测试文件是否可访问
                        file_path.stat()
                        files.append(file_path)
                        print(f"找到文件: {file_path.name}")
                    except (OSError, UnicodeError) as e:
                        print(f"跳过无法访问的文件 {file_path.name}: {e}")
                        continue
            
            # 改进文件排序逻辑
            def sort_key(file_path):
                filename = file_path.name
                # 提取数字前缀
                match = re.match(r'^(\d+)', filename)
                if match:
                    return (0, int(match.group(1)), filename)  # 数字开头的文件优先
                else:
                    return (1, filename)  # 非数字开头的文件按字母顺序排序
            
            files.sort(key=sort_key)
            print(f"文件排序后: {[f.name for f in files]}")
            print(f"总共找到 {len(files)} 个文件")
            return files
        except Exception as e:
            print(f"扫描学生文件夹失败: {e}")
            import traceback
            traceback.print_exc()
            return files

    def merge_pdfs_with_bookmarks(self, pdf_files, output_path, student_name):
        try:
            print(f"\n开始合并PDF，共 {len(pdf_files)} 个文件")
            merged_doc = fitz.open()
            bookmarks = []
            temp_files = []
            
            for i, file_path in enumerate(pdf_files):
                print(f"\n处理文件 {i+1}/{len(pdf_files)}: {file_path.name}")
                current_pdf_path = None
                
                if file_path.suffix.lower() == '.pdf':
                    current_pdf_path = file_path
                    print(f"  直接使用PDF文件: {current_pdf_path}")
                elif file_path.suffix.lower() in ['.doc', '.docx']:
                    print(f"  转换Word文件为PDF: {file_path.name}")
                    temp_pdf_path = self.word_to_pdf(file_path, file_path.parent)
                    if temp_pdf_path and temp_pdf_path.exists():
                        current_pdf_path = temp_pdf_path
                        temp_files.append(temp_pdf_path)
                        print(f"  Word转换成功: {temp_pdf_path}")
                    else:
                        print(f"  跳过文件（转换失败）: {file_path.name}")
                        continue
                else:
                    print(f"  跳过不支持的文件类型: {file_path.name}")
                    continue
                
                if current_pdf_path:
                    try:
                        print(f"  正在合并PDF: {current_pdf_path}")
                        current_doc = fitz.open(str(current_pdf_path))
                        current_page_count = merged_doc.page_count
                        print(f"  当前合并文档页数: {current_page_count}")
                        print(f"  要添加的文档页数: {current_doc.page_count}")
                        
                        merged_doc.insert_pdf(current_doc)
                        
                        # 确保书签标题正确处理中文编码
                        bookmark_title = self.ensure_utf8_string(file_path.stem)
                        bookmarks.append([1, bookmark_title, current_page_count + 1])
                        current_doc.close()
                        print(f"  ✓ 已添加: {bookmark_title} (从第 {current_page_count + 1} 页开始)")
                    except Exception as e:
                        print(f"  ✗ 处理PDF文件失败 {current_pdf_path}: {e}")
                        continue
            
            print(f"\n合并完成，总共 {merged_doc.page_count} 页")
            print(f"书签数量: {len(bookmarks)}")
            
            # 设置PDF元数据以支持中文
            merged_doc.set_metadata({
                "title": f"{student_name}的论文材料",
                "author": "系统生成",
                "subject": "学生论文材料合并",
                "creator": "StudentDocumentMerger",
                "producer": "PyMuPDF"
            })
            
            # 添加书签目录，确保中文正确显示
            if bookmarks:
                try:
                    # 使用set_toc方法添加书签，确保UTF-8编码
                    merged_doc.set_toc(bookmarks)
                    print(f"成功添加 {len(bookmarks)} 个书签")
                except Exception as e:
                    print(f"添加书签失败: {e}")
                    # 如果set_toc失败，尝试使用底层API添加书签
                    try:
                        # 清空现有书签
                        merged_doc.set_toc([])
                        # 逐个添加书签
                        for level, title, page in bookmarks:
                            # 确保标题是UTF-8字符串
                            utf8_title = self.ensure_utf8_string(title)
                            # 使用底层API添加书签
                            merged_doc.set_toc(merged_doc.get_toc() + [(level, utf8_title, page)])
                        print(f"使用底层API成功添加 {len(bookmarks)} 个书签")
                    except Exception as e2:
                        print(f"底层API添加书签也失败: {e2}")
                        # 最后尝试：直接设置整个目录
                        try:
                            toc = []
                            for level, title, page in bookmarks:
                                utf8_title = self.ensure_utf8_string(title)
                                toc.append((level, utf8_title, page))
                            merged_doc.set_toc(toc)
                            print(f"直接设置目录成功，添加 {len(toc)} 个书签")
                        except Exception as e3:
                            print(f"所有书签添加方法都失败: {e3}")
            
            # 确保输出路径使用正确的编码
            output_path_str = str(output_path)
            print(f"保存PDF到: {output_path_str}")
            
            # 确保输出目录存在
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            merged_doc.save(output_path_str)
            merged_doc.close()
            
            # 验证文件是否成功保存
            if Path(output_path_str).exists():
                print(f"PDF文件保存成功: {output_path_str}")
            else:
                print(f"警告: PDF文件保存后不存在: {output_path_str}")
            
            # 清理临时文件
            for temp_file in temp_files:
                try:
                    temp_file.unlink()
                except Exception:
                    pass
            print(f"合并PDF已保存: {output_path}")
            return True
        except Exception as e:
            print(f"合并PDF失败: {e}")
            return False

    def process_student(self, student_folder):
        try:
            # 安全地获取学生名称
            student_path = Path(student_folder)
            student_name = student_path.name
            
            # 确保学生名称正确编码
            try:
                # 尝试修复可能的编码问题
                if isinstance(student_name, bytes):
                    student_name = student_name.decode('utf-8')
                elif isinstance(student_name, str):
                    # 检查是否包含乱码
                    if any(ord(c) > 127 and ord(c) < 256 for c in student_name):
                        # 可能是latin1编码的中文，尝试转换
                        try:
                            student_name = student_name.encode('latin1').decode('utf-8')
                        except:
                            pass
            except Exception as e:
                print(f"处理学生名称编码时出错: {e}")
            
            # 清理学生名称
            safe_student_name = self.sanitize_filename(student_name)
            
            print(f"\n{'='*50}")
            print(f"开始处理学生: {student_name}")
            print(f"安全学生名称: {safe_student_name}")
            print(f"学生文件夹: {student_folder}")
            print(f"文件夹路径存在: {student_path.exists()}")
            
            # 检查学生文件夹是否存在
            if not student_path.exists():
                print(f"错误: 学生文件夹不存在: {student_folder}")
                return False
            
            # 创建封面
            cover_path = self.update_cover_for_student(student_name)
            if cover_path:
                print(f"封面创建成功: {cover_path}")
                print(f"封面文件是否存在: {cover_path.exists()}")
                if cover_path.exists():
                    print(f"封面文件大小: {cover_path.stat().st_size} 字节")
                else:
                    print("警告: 封面文件创建后不存在!")
            else:
                print(f"警告: 无法为学生 {student_name} 创建封面")
            
            # 获取学生文件
            files = self.get_student_files(student_folder)
            print(f"找到 {len(files)} 个文件: {[f.name for f in files]}")
            
            pdf_files = []
            temp_files_to_delete = []
            
            # 首先处理封面PDF（确保封面在最前面）
            if cover_path:
                print(f"\n🔄 正在转换封面为PDF...")
                print(f"  封面文件: {cover_path}")
                print(f"  封面文件大小: {cover_path.stat().st_size} 字节")
                print(f"  输出目录: {cover_path.parent}")
                
                cover_pdf_path = self.word_to_pdf(cover_path, cover_path.parent)
                if cover_pdf_path and cover_pdf_path.exists():
                    pdf_files.append(cover_pdf_path)
                    temp_files_to_delete.append(cover_pdf_path)
                    print(f"  ✓ 封面PDF转换成功: {cover_pdf_path}")
                    print(f"  封面PDF大小: {cover_pdf_path.stat().st_size} 字节")
                else:
                    print(f"  ✗ 封面PDF转换失败")
                    print(f"  请检查封面模板文件是否损坏或LibreOffice是否正确安装")
            else:
                print("⚠️ 警告: 没有封面文件")
            
            # 处理其他文件
            print("正在处理其他文件...")
            for f in files:
                print(f"\n处理文件: {f.name} (类型: {f.suffix})")
                if f.suffix.lower() == '.pdf':
                    pdf_files.append(f)
                    print(f"  ✓ 直接添加PDF文件: {f.name}")
                elif f.suffix.lower() in ['.doc', '.docx']:
                    print(f"  🔄 正在转换Word文件: {f.name}")
                    print(f"    输入文件大小: {f.stat().st_size} 字节")
                    pdf_path = self.word_to_pdf(f, f.parent)
                    if pdf_path and pdf_path.exists():
                        pdf_files.append(pdf_path)
                        temp_files_to_delete.append(pdf_path)
                        print(f"  ✓ Word文件转换成功: {f.name} -> {pdf_path.name}")
                        print(f"    输出PDF大小: {pdf_path.stat().st_size} 字节")
                    else:
                        print(f"  ✗ Word文件转换失败: {f.name}")
                        print(f"    请检查文件是否损坏或LibreOffice是否正确安装")
                else:
                    print(f"  ⏭️ 跳过不支持的文件类型: {f.name}")
            
            print(f"总共收集到 {len(pdf_files)} 个PDF文件")
            print(f"PDF文件列表:")
            for i, pdf_file in enumerate(pdf_files):
                print(f"  {i+1}. {pdf_file.name} ({pdf_file})")
            
            # 合并PDF - 确保路径正确处理
            safe_student_name = self.sanitize_filename(student_name)
            output_filename = f"{safe_student_name}.pdf"
            output_pdf = self.safe_path_join(self.output_dir, output_filename)
            print(f"正在合并PDF到: {output_pdf}")
            print(f"输出目录: {self.output_dir}")
            print(f"输出文件名: {output_filename}")
            
            success = self.merge_pdfs_with_bookmarks(pdf_files, output_pdf, student_name)
            if success:
                print(f"学生 {student_name} 处理完成: {output_pdf}")
            else:
                print(f"学生 {student_name} 处理失败")
            
            # 清理临时文件
            if cover_path and cover_path.exists():
                try:
                    cover_path.unlink()
                    print(f"已删除临时封面文件: {cover_path}")
                except Exception as e:
                    print(f"删除临时封面文件失败: {e}")
            
            for temp_file in temp_files_to_delete:
                if temp_file.exists():
                    try:
                        temp_file.unlink()
                        print(f"已删除临时PDF文件: {temp_file}")
                    except Exception as e:
                        print(f"删除临时PDF文件失败: {e}")
            
            return success
            
        except Exception as e:
            print(f"处理学生 {student_name} 时发生错误: {e}")
            import traceback
            traceback.print_exc()
            return False

    def run(self):
        print(f"\n{'='*60}")
        print(f"开始处理重点审议材料合并")
        print(f"学生文件夹数量: {len(self.student_folders)}")
        print(f"输出目录: {self.output_dir}")
        print(f"Excel文件: {self.excel_file}")
        print(f"封面模板: {self.cover_template}")
        print(f"{'='*60}")
        
        success_count = 0
        failed_count = 0
        
        for i, folder in enumerate(self.student_folders, 1):
            print(f"\n进度: {i}/{len(self.student_folders)}")
            success = self.process_student(folder)
            if success:
                success_count += 1
            else:
                failed_count += 1
        
        print(f"\n{'='*60}")
        print(f"处理完成!")
        print(f"成功处理: {success_count} 个学生")
        print(f"处理失败: {failed_count} 个学生")
        print(f"总计: {len(self.student_folders)} 个学生")
        print(f"{'='*60}")
        
        return success_count > 0

def main():
    if len(sys.argv) < 2:
        print("请传入JSON参数，包含student_folders和output_dir")
        sys.exit(1)
    try:
        # 输出系统编码信息
        print(f"系统默认编码: {sys.getdefaultencoding()}")
        print(f"文件系统编码: {sys.getfilesystemencoding()}")
        print(f"当前工作目录: {os.getcwd()}")
        print(f"Python版本: {sys.version}")
        
        args = json.loads(sys.argv[1])
        student_folders = args.get('student_folders', [])
        output_dir = args.get('output_dir', './合并后的PDF文件')
        excel_file = args.get('excel_file')
        base_dir = Path(__file__).parent
        
        print(f"接收到的参数:")
        print(f"  学生文件夹: {student_folders}")
        print(f"  输出目录: {output_dir}")
        print(f"  Excel文件: {excel_file}")
        print(f"  基础目录: {base_dir}")
        
        # 验证所有路径
        for i, folder in enumerate(student_folders):
            folder_path = Path(folder)
            print(f"  学生文件夹 {i+1}: {folder}")
            print(f"    路径存在: {folder_path.exists()}")
            print(f"    是否为目录: {folder_path.is_dir() if folder_path.exists() else 'N/A'}")
            print(f"    路径编码: {repr(str(folder_path))}")
        
        merger = StudentDocumentMerger(student_folders, output_dir, base_dir, excel_file=excel_file)
        success = merger.run()
        print(json.dumps({"success": success}))
    except Exception as e:
        print(f"主程序执行失败: {e}")
        import traceback
        traceback.print_exc()
        print(json.dumps({"success": False, "error": str(e)}))
        sys.exit(1)

if __name__ == "__main__":
    main() 